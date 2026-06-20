"""
docx_builder.py
Converts a Document model instance to a formatted .docx file using python-docx.
Canonical style: Times New Roman 12pt headings / 10pt body, no section rules,
page numbers from page 2, header trimmed to agreement type only.
"""
import io
import re

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup, NavigableString


def _html_to_markdown(html: str) -> str:
    """Convert HTML (from Quill/CKEditor) to the markdown-like format the docx builder understands."""
    if "<" not in html or ">" not in html:
        return html

    soup = BeautifulSoup(html, "html.parser")
    lines = []

    def inline(node):
        parts = []
        for child in node.children:
            if isinstance(child, NavigableString):
                parts.append(str(child))
            elif child.name in ("strong", "b"):
                parts.append(f"**{inline(child)}**")
            elif child.name in ("em", "i"):
                parts.append(f"*{inline(child)}*")
            elif child.name == "u":
                parts.append(inline(child))
            elif child.name == "br":
                parts.append("\n")
            elif child.name == "a":
                parts.append(inline(child))
            else:
                parts.append(inline(child))
        return "".join(parts)

    for el in soup.children:
        if isinstance(el, NavigableString):
            txt = str(el).strip()
            if txt:
                lines.append(txt)
            continue
        name = el.name
        if name == "h1":
            lines.append(f"# {inline(el).strip()}")
            lines.append("")
        elif name == "h2":
            lines.append(f"## {inline(el).strip()}")
            lines.append("")
        elif name == "h3":
            lines.append(f"## {inline(el).strip()}")
            lines.append("")
        elif name == "p":
            txt = inline(el).strip()
            if txt:
                lines.append(txt)
            lines.append("")
        elif name == "ul":
            for li in el.find_all("li", recursive=False):
                lines.append(f"• {inline(li).strip()}")
            lines.append("")
        elif name == "ol":
            for i, li in enumerate(el.find_all("li", recursive=False), start=1):
                lines.append(f"{i}. {inline(li).strip()}")
            lines.append("")
        elif name == "hr":
            lines.append("---")
            lines.append("")
        elif name == "blockquote":
            lines.append(inline(el).strip())
            lines.append("")
        else:
            txt = inline(el).strip()
            if txt:
                lines.append(txt)
                lines.append("")

    return "\n".join(lines)


# ── Colour constants ────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1A, 0x2E, 0x4A)
GRAY  = RGBColor(0x88, 0x88, 0x88)
RED   = RGBColor(0xCC, 0x00, 0x00)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
ANCHOR_GRAY = RGBColor(0xCC, 0xCC, 0xCC)  # light gray for DocuSign anchors

# DocuSign anchor tokens like /sign1/ or /sign_celebrity/.
# Rendered in tiny light-gray text so they are present in the PDF text layer
# and reliably found by DocuSign anchor scanner. DocuSign tab overlays
# (signature stamps, name fields, etc.) cover this text in the signed PDF.
_DS_ANCHOR_RE = re.compile(r"(/[a-z][a-z0-9_]*/)")


def _set_font(run, name: str, size_pt: float, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _reorder_signature_before_exhibits(content: str) -> str:
    """If a SIGNATURE PAGE section appears after EXHIBIT sections, move it before them."""
    sig_idx = content.find("SIGNATURE PAGE")
    if sig_idx == -1:
        sig_idx = content.find("IN WITNESS WHEREOF")
    exhibit_idx = content.find("EXHIBIT A")
    if exhibit_idx == -1:
        exhibit_idx = content.find("EXHIBIT B")

    if sig_idx == -1 or exhibit_idx == -1 or sig_idx < exhibit_idx:
        return content

    sig_section_start = content.rfind("\n", 0, sig_idx)
    if sig_section_start == -1:
        sig_section_start = 0
    heading_match = re.search(r"(\n#+\s*SIGNATURE[^\n]*\n)", content[:sig_idx + 50])
    if heading_match:
        sig_section_start = heading_match.start()

    sig_block = content[sig_section_start:].strip()
    pre_sig = content[:sig_section_start]
    exhibit_start = pre_sig.find("EXHIBIT")
    if exhibit_start == -1:
        return content
    exhibit_heading = re.search(r"(\n#+\s*EXHIBIT[^\n]*\n)", pre_sig[:exhibit_start + 50])
    if exhibit_heading:
        exhibit_section_start = exhibit_heading.start()
    else:
        exhibit_section_start = pre_sig.rfind("\n", 0, exhibit_start)
    main_body = pre_sig[:exhibit_section_start].rstrip()
    exhibits_block = pre_sig[exhibit_section_start:].strip()
    return f"{main_body}\n\n{sig_block}\n\n{exhibits_block}\n"


def _parse_inline(paragraph, text: str, base_size: float, base_color=None):
    """Parse **bold**, *italic*, and /anchor/ markers; anchors render invisible white."""
    # Split on anchor tokens first so they survive the bold/italic pass
    segments = _DS_ANCHOR_RE.split(text)
    for seg in segments:
        if _DS_ANCHOR_RE.fullmatch(seg):
            run = paragraph.add_run(seg)
            _set_font(run, "Times New Roman", 6, color=ANCHOR_GRAY)
            continue
        # Normal text — split on bold/italic markers
        bold_italic = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
        for part in bold_italic.split(seg):
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                _set_font(run, "Times New Roman", base_size, bold=True, color=base_color)
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(part[1:-1])
                _set_font(run, "Times New Roman", base_size, italic=True, color=base_color)
            elif part:
                run = paragraph.add_run(part)
                _set_font(run, "Times New Roman", base_size, color=base_color)


def _set_margins(doc, inches=1.0):
    for section in doc.sections:
        section.top_margin    = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin   = Inches(inches)
        section.right_margin  = Inches(inches)


def _add_header(doc, title: str):
    """Running-page header. Show only the agreement type — strip everything after
    the first em-dash or en-dash so long contracting-party strings don't appear."""
    short = (title or "Agreement").split("—")[0].split("–")[0].strip()
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    run = hp.add_run(short.upper())
    _set_font(run, "Times New Roman", 8, color=GRAY)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_footer(doc):
    """Footer: page number right only. Page 1 has no page number."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    sectPr = section._sectPr
    if sectPr.find(qn("w:titlePg")) is None:
        sectPr.append(OxmlElement("w:titlePg"))

    # First-page footer — blank
    fp_first = section.first_page_footer.paragraphs[0] if section.first_page_footer.paragraphs else section.first_page_footer.add_paragraph()
    fp_first.clear()

    # Subsequent-pages footer: page number right-aligned only
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()

    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    fp.add_run("\t")

    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")

    page_run = fp.add_run()
    page_run.font.name = "Times New Roman"
    page_run.font.size = Pt(8)
    page_run.font.color.rgb = GRAY
    r_elem = page_run._r
    r_elem.append(fld_char1)
    r_elem.append(instr)
    r_elem.append(fld_char2)
    r_elem.append(fld_char3)


def build_docx_with_track_changes(doc_model) -> bytes:
    """Same as build_docx but opens in Word with Track Changes already ON."""
    raw = build_docx(doc_model)
    import zipfile, tempfile
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as src:
        src.write(raw); src.flush()
        src_path = src.name
    out_buf = io.BytesIO()
    with zipfile.ZipFile(src_path, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.namelist():
            data = zin.read(item)
            if item == "word/settings.xml":
                txt = data.decode("utf-8")
                if "<w:trackChanges" not in txt:
                    txt = re.sub(r"(<w:settings\b[^>]*>)", r'\1<w:trackChanges/>', txt, count=1)
                data = txt.encode("utf-8")
            zout.writestr(item, data)
    return out_buf.getvalue()


def build_docx(doc_model) -> bytes:
    """Convert a Document model instance to .docx bytes.
    Canonical style: Times New Roman 12pt headings / 10pt body, no section rules,
    page numbers from page 2, header shows agreement type only."""
    docx = DocxDocument()
    _set_margins(docx, 1.0)
    _add_header(docx, doc_model.title)
    _add_footer(docx)

    if getattr(doc_model, "status", "draft") in ("draft", "negotiating"):
        banner = docx.add_paragraph()
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = banner.add_run("DRAFT — ATTORNEY REVIEW REQUIRED")
        _set_font(run, "Times New Roman", 9, bold=True, italic=True, color=RED)
        banner.paragraph_format.space_after = Pt(6)

    content = doc_model.content or ""
    if "<p>" in content or "<h1>" in content or "<h2>" in content or "<strong>" in content:
        content = _html_to_markdown(content)
    content = _reorder_signature_before_exhibits(content)
    lines = content.splitlines()
    prev_blank = False

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # HTML <table> block
        if stripped.lower().startswith("<table"):
            tbl_buf = [line]
            j = i + 1
            while j < len(lines):
                tbl_buf.append(lines[j])
                if "</table>" in lines[j].lower():
                    j += 1
                    break
                j += 1
            try:
                soup = BeautifulSoup("\n".join(tbl_buf), "html.parser")
                html_tbl = soup.find("table")
                rows_html = []
                for tr in html_tbl.find_all("tr"):
                    cells = [c.get_text(" ", strip=True) for c in tr.find_all(["th", "td"])]
                    if cells:
                        rows_html.append(cells)
                if rows_html:
                    ncols = max(len(r) for r in rows_html)
                    tbl = docx.add_table(rows=len(rows_html), cols=ncols)
                    tbl.style = "Light Grid Accent 1"
                    has_header = bool(html_tbl.find("thead")) or bool(html_tbl.find("th"))
                    for r_idx, row in enumerate(rows_html):
                        for c_idx, cell_text in enumerate(row):
                            cell = tbl.rows[r_idx].cells[c_idx]
                            cell.text = ""
                            p = cell.paragraphs[0]
                            _parse_inline(p, cell_text, 10, base_color=NAVY if (has_header and r_idx == 0) else None)
                            if has_header and r_idx == 0:
                                for run in p.runs:
                                    run.font.bold = True
            except Exception:
                for raw in tbl_buf:
                    docx.add_paragraph(raw)
            i = j
            prev_blank = False
            continue

        # Markdown pipe-table
        if stripped.startswith("|") and stripped.endswith("|") and (i + 1) < len(lines):
            sep = lines[i + 1].strip()
            if re.match(r"^\|[\s\-:|]+\|$", sep):
                table_lines = [stripped]
                j = i + 2
                while j < len(lines) and lines[j].strip().startswith("|") and lines[j].strip().endswith("|"):
                    table_lines.append(lines[j].strip())
                    j += 1
                rows = [[c.strip() for c in row.strip("|").split("|")] for row in table_lines]
                if rows:
                    ncols = len(rows[0])
                    tbl = docx.add_table(rows=len(rows), cols=ncols)
                    tbl.style = "Light Grid Accent 1"
                    for r_idx, row in enumerate(rows):
                        for c_idx, cell_text in enumerate(row):
                            if c_idx >= ncols:
                                break
                            cell = tbl.rows[r_idx].cells[c_idx]
                            cell.text = ""
                            p = cell.paragraphs[0]
                            _parse_inline(p, cell_text, 10, base_color=NAVY if r_idx == 0 else None)
                            for run in p.runs:
                                if r_idx == 0:
                                    run.font.bold = True
                i = j
                prev_blank = False
                continue

        i += 1

        # Horizontal rule — skipped. Section spacing is controlled by
        # heading/paragraph space_before/space_after, not separator lines.
        if stripped == "---":
            prev_blank = False
            continue

        # Blank line — skipped. Same rationale as above.
        if stripped == "":
            prev_blank = True
            continue
        prev_blank = False

        # Standalone DocuSign anchor line (e.g. /sign_celebrity/) — rendered as
        # invisible white text so DocuSign can find it but readers don't see it.
        if _DS_ANCHOR_RE.fullmatch(stripped):
            p = docx.add_paragraph()
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(stripped)
            _set_font(run, "Times New Roman", 6, color=ANCHOR_GRAY)
            continue

        # Title: # HEADING — TNR 12 bold, centered
        if re.match(r"^#\s+", stripped):
            text = re.sub(r"^#+\s+", "", stripped)
            if re.match(r"^(SIGNATURE PAGE|EXHIBIT\s+[A-Z])", text, re.I):
                _add_page_break(docx)
            p = docx.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(6)
            run = p.add_run(text)
            _set_font(run, "Times New Roman", 12, bold=True)
            continue

        # Section heading: ## HEADING — TNR 12 bold, left
        if re.match(r"^##\s+", stripped):
            text = re.sub(r"^#+\s+", "", stripped)
            if re.match(r"^(SIGNATURE PAGE|EXHIBIT\s+[A-Z])", text, re.I):
                _add_page_break(docx)
            p = docx.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(text)
            _set_font(run, "Times New Roman", 12, bold=True)
            continue

        # Sub-heading ### or deeper — TNR 12 bold italic
        if re.match(r"^#{3,}\s+", stripped):
            text = re.sub(r"^#+\s+", "", stripped)
            p = docx.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(text)
            _set_font(run, "Times New Roman", 12, bold=True, italic=True)
            continue

        # Normal paragraph — TNR 10 body, tight spacing
        p = docx.add_paragraph()
        p.paragraph_format.line_spacing = 1.15

        # Signature-block field lines get extra breathing room so they
        # don't run into each other in the signed PDF.
        _SIG_FIELD_RE = re.compile(
            r'^(Signature|By|Print Name|Title|Date|Address|Phone|Email|'
            r'SOC\. SEC\.|SS#|In Case of Emergency|Emergency Phone|'
            r'EMPLOYEE:|PRODUCER:|EMPLOYER:)',
            re.IGNORECASE
        )
        if _SIG_FIELD_RE.match(stripped):
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(4)
        else:
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.space_before = Pt(0)

        _parse_inline(p, stripped, 10)

    buf = io.BytesIO()
    docx.save(buf)
    buf.seek(0)
    return buf.getvalue()
